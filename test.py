import os
import io
import streamlit as st
import whisper
import os
import sounddevice as sd
import numpy as np
from scipy.io.wavfile import write
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import fitz  # PyMuPDF
import docx
from openai import OpenAI

# Load API key from environment variable
os.environ["OPENAI_API_KEY"] = "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"
client = OpenAI()


nltk.download("punkt", quiet=True)
nltk.download("punkt_tab", quiet=True)
nltk.download("stopwords", quiet=True)


def extract_text_from_pdf(file):
    if hasattr(file, "read"):
        file_bytes = io.BytesIO(file.read())
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            return "".join([page.get_text() for page in doc])
    else:
        with fitz.open(file) as doc:
            return "".join([page.get_text() for page in doc])


def extract_text_from_docx(file):
    if hasattr(file, "read"):
        file_bytes = io.BytesIO(file.read())
        doc = docx.Document(file_bytes)
    else:
        doc = docx.Document(file)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_keywords(text, top_n=15):
    tokens = word_tokenize(text.lower())
    words = [w for w in tokens if w.isalpha() and w not in stopwords.words("english")]
    freq_dist = nltk.FreqDist(words)
    return [word for word, _ in freq_dist.most_common(top_n)]


def record_audio(duration=5, filename="question.wav"):
    fs = 16000
    st.info(f"🎙 Recording for {duration} seconds...")
    audio = sd.rec(int(duration * fs), samplerate=fs, channels=1)
    sd.wait()
    write(filename, fs, (audio * 32767).astype(np.int16))
    return filename


def transcribe_audio(filename="question.wav"):
    model = whisper.load_model("base")
    result = model.transcribe(filename)
    return result["text"]


def generate_star_answer(question, resume_text, job_keywords):
    prompt = f"""
Answer this interview question in STAR (Situation, Task, Action, Result) format:
Question: {question}

Candidate Resume:
{resume_text}

Job Keywords: {', '.join(job_keywords)}

Keep it professional, confident, and concise.
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7
    )
    return response.choices[0].message.content

def run_streamlit():
    st.set_page_config(page_title="AI Interview Copilot", page_icon="🎙", layout="centered")
    st.title("🎙 AI Interview Copilot")

    resume_file = st.file_uploader("📄 Upload Resume (PDF or DOCX)", type=["pdf", "docx"])
    jd_file = st.file_uploader("📜 Upload Job Description (PDF or DOCX)", type=["pdf", "docx"])

    if resume_file and jd_file:
        resume_text = extract_text_from_pdf(resume_file) if resume_file.name.endswith(".pdf") else extract_text_from_docx(resume_file)
        jd_text = extract_text_from_pdf(jd_file) if jd_file.name.endswith(".pdf") else extract_text_from_docx(jd_file)
        job_keywords = extract_keywords(jd_text)

        st.subheader("📄 Resume Context")
        st.text_area("Resume", resume_text, height=200)
        st.subheader("🔑 Job Description Keywords")
        st.write(", ".join(job_keywords))

        if st.button("🎤 Record Interview Question"):
            audio_file = record_audio()
            st.success(f"✅ Audio recorded: {audio_file}")

            question_text = transcribe_audio(audio_file)
            st.subheader("🎙 Transcribed Question")
            st.write(question_text)

            answer = generate_star_answer(question_text, resume_text, job_keywords)
            st.subheader("🤖 AI STAR Answer")
            st.write(answer)

if __name__ == "__main__":
    run_streamlit()
